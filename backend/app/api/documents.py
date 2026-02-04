"""Document upload and management API routes."""
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from sqlalchemy.orm import Session
from app.models.base import get_db
from app.models.document import Document, DocumentChunk
from app.services.tools import KnowledgeBaseTool
from app.services.ai_assistant import AIAssistant
import uuid
import os
from typing import List
import PyPDF2
from docx import Document as DocxDocument

router = APIRouter(prefix="/api/documents", tags=["documents"])


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file."""
    try:
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text
    except Exception as e:
        raise Exception(f"Error extracting PDF text: {str(e)}")


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file."""
    try:
        doc = DocxDocument(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        raise Exception(f"Error extracting DOCX text: {str(e)}")


def chunk_text(text: str, chunk_size: int = 500) -> List[dict]:
    """Split text into chunks for embedding."""
    words = text.split()
    chunks = []
    
    current_chunk = []
    current_length = 0
    
    for word in words:
        current_chunk.append(word)
        current_length += len(word) + 1  # +1 for space
        
        if current_length >= chunk_size:
            chunk_text = " ".join(current_chunk)
            chunks.append({
                "text": chunk_text,
                "chunk_index": len(chunks)
            })
            current_chunk = []
            current_length = 0
    
    # Add remaining text as final chunk
    if current_chunk:
        chunk_text = " ".join(current_chunk)
        chunks.append({
            "text": chunk_text,
            "chunk_index": len(chunks)
        })
    
    return chunks


@router.post("/upload")
async def upload_document(
    user_id: str,
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    """Upload and process a document."""
    try:
        # Create document record
        document = Document(
            user_id=uuid.UUID(user_id),
            filename=file.filename,
            file_type=file.filename.split('.')[-1].lower(),
            file_size=0,  # Will update after saving
            status="processing"
        )
        db.add(document)
        db.commit()
        db.refresh(document)
        
        # Save file temporarily
        upload_dir = "uploads"
        os.makedirs(upload_dir, exist_ok=True)
        file_path = os.path.join(upload_dir, f"{document.id}_{file.filename}")
        
        with open(file_path, "wb") as buffer:
            content = await file.read()
            buffer.write(content)
            document.file_size = len(content)
        
        # Extract text based on file type
        if document.file_type == "pdf":
            text = extract_text_from_pdf(file_path)
        elif document.file_type in ["docx", "doc"]:
            text = extract_text_from_docx(file_path)
        elif document.file_type == "txt":
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {document.file_type}")
        
        # Chunk the text
        chunks = chunk_text(text)
        
        # Store chunks in database
        for chunk_data in chunks:
            chunk = DocumentChunk(
                document_id=document.id,
                chunk_text=chunk_data["text"],
                chunk_index=chunk_data["chunk_index"]
            )
            db.add(chunk)
        
        db.commit()
        
        # Add to vector database
        knowledge_base = KnowledgeBaseTool()
        chunks_with_source = [
            {
                "text": chunk_data["text"],
                "chunk_index": chunk_data["chunk_index"],
                "source": file.filename
            }
            for chunk_data in chunks
        ]
        
        success = knowledge_base.add_document_chunks(
            user_id=user_id,
            document_id=str(document.id),
            chunks=chunks_with_source
        )
        
        if success:
            document.status = "completed"
        else:
            document.status = "failed"
        
        db.commit()
        
        # Clean up temp file
        os.remove(file_path)
        
        return {
            "document_id": str(document.id),
            "filename": document.filename,
            "status": document.status,
            "chunks": len(chunks)
        }
        
    except Exception as e:
        if document:
            document.status = "failed"
            db.commit()
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/{user_id}")
async def get_documents(
    user_id: str,
    db: Session = Depends(get_db)
):
    """Get all documents for a user."""
    try:
        documents = db.query(Document).filter(
            Document.user_id == uuid.UUID(user_id)
        ).order_by(Document.created_at.desc()).all()
        
        return [
            {
                "id": str(doc.id),
                "filename": doc.filename,
                "file_type": doc.file_type,
                "file_size": doc.file_size,
                "status": doc.status,
                "created_at": doc.created_at.isoformat(),
                "chunk_count": len(doc.chunks)
            }
            for doc in documents
        ]
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

